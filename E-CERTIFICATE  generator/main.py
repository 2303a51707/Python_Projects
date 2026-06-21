from docx import Document
import os

placeholder = "[name]"
template_path = "D:/certificate.docx"
names_path = "D:/data.txt"
output_dir = "D:/certificates"

os.makedirs(output_dir, exist_ok=True)

with open(names_path) as names_file:
    names = [name.strip() for name in names_file]

for name in names:
    doc = Document(template_path)

    for para in doc.paragraphs:
        if placeholder in para.text:
            for run in para.runs:
                run.text = run.text.replace(placeholder, name)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = run.text.replace(placeholder, name)

    for section in doc.sections:
        for header_para in section.header.paragraphs:
            for run in header_para.runs:
                run.text = run.text.replace(placeholder, name)

        for footer_para in section.footer.paragraphs:
            for run in footer_para.runs:
                run.text = run.text.replace(placeholder, name)

    output_path = os.path.join(output_dir, f"{name}.docx")
    doc.save(output_path)
    print(f"Certificate created for {name}")